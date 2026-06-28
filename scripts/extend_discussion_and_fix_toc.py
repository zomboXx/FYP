from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path("docs/find-your-path-report.docx")


def style_name(paragraph):
    return paragraph.style.name if paragraph.style is not None else ""


def set_font(run, bold=None):
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    rpr.get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)
    if bold is not None:
        run.bold = bold


def format_para(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    for run in paragraph.runs:
        set_font(run)


def heading(doc, text):
    matches = [p for p in doc.paragraphs if style_name(p).startswith("Heading") and p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one heading {text!r}, found {len(matches)}")
    return matches[0]


def add_heading_before(anchor, text):
    p = anchor.insert_paragraph_before("", style="Heading 2")
    r = p.add_run(text)
    set_font(r, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para_before(anchor, text):
    p = anchor.insert_paragraph_before(text)
    format_para(p)
    return p


def remove_static_toc(doc):
    # Keep the "MỤC LỤC" title and delete static TOC paragraphs that follow it.
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "MỤC LỤC":
            start = i
            break
    if start is None:
        return
    for p in list(doc.paragraphs[start + 1:]):
        if style_name(p).startswith("toc"):
            p._p.getparent().remove(p._p)
        elif p.text.strip().startswith("I.") or style_name(p).startswith("Heading"):
            break


def main():
    doc = Document(DOCX)
    conclusion = heading(doc, "V. Kết luận")

    sections = [
        ("4. Kịch bản vận hành đề xuất cho Find Your Path", [
            "Một cách triển khai hợp lý cho Find Your Path là xem hệ thống như chuỗi nhiều tầng quyết định. Tầng đầu tiên nhận yêu cầu giao hàng, chuẩn hóa vị trí, xác định điểm xuất phát và điểm đích. Tầng thứ hai dùng A* để tính tuyến đường cơ sở trên graph hiện biết. Tầng thứ ba kiểm tra các ràng buộc như deadline, capacity và thứ tự phục vụ bằng Forward Checking. Nếu route hợp lệ, hệ thống có thể trả kết quả; nếu chưa hợp lệ, hệ thống cần điều chỉnh cấu hình hoặc thử phương án khác.",
            "Trong kịch bản nhiều đơn hàng, Simple Hill Climbing có thể được đặt trước hoặc sau A*. Nếu đặt trước, thuật toán tạo thứ tự điểm đến sơ bộ, sau đó A* tính chi phí giữa từng cặp điểm. Nếu đặt sau, thuật toán nhận một route ban đầu và thử cải thiện bằng các phép hoán đổi. Cả hai cách đều hợp lý, nhưng báo cáo nên nêu rõ cách nhóm chọn để tránh hiểu nhầm rằng Simple Hill Climbing đang thay thế hoàn toàn A*.",
            "AND-OR Search có thể được dùng ở tầng mô phỏng môi trường bất định. Ví dụ, khi hệ thống biết một đoạn đường có xác suất bị chặn, nó có thể tạo kế hoạch gồm nhánh chính và nhánh dự phòng. Nếu đường thông, shipper đi theo route ngắn; nếu đường bị chặn, shipper chuyển sang route thay thế. Cách trình bày này giúp AND-OR Search có vai trò rõ ràng hơn là chỉ một thuật toán lý thuyết.",
            "Alpha-Beta Pruning có thể nằm ở tầng đánh giá rủi ro. Khi có nhiều route ứng viên, hệ thống có thể mô phỏng các phản ứng bất lợi của môi trường và chọn route có utility tốt hơn trong trường hợp xấu. Điều này không biến bài toán giao hàng thành trò chơi hoàn chỉnh, nhưng giúp giải thích tại sao tư duy đối kháng vẫn có ý nghĩa trong bài toán vận tải.",
            "Kịch bản vận hành này cũng cho thấy lý do báo cáo không nên dùng bảng đánh giá đơn giản. Mỗi thuật toán có vị trí khác nhau trong pipeline, nên so sánh trực tiếp theo một thước đo duy nhất sẽ làm nghèo phân tích. Thay vào đó, cần mô tả quan hệ phối hợp: thuật toán nào tạo route, thuật toán nào kiểm tra ràng buộc, thuật toán nào lập kế hoạch dự phòng và thuật toán nào đánh giá rủi ro.",
            "Nếu hiện tại project chưa triển khai đầy đủ toàn bộ pipeline, báo cáo vẫn có thể mô tả kịch bản vận hành như định hướng thiết kế. Điều này không làm sai nội dung, miễn là nhóm phân biệt rõ phần đã cài đặt, phần đang mô phỏng và phần đề xuất mở rộng. Trong báo cáo học phần, khả năng phân tích kiến trúc thuật toán cũng quan trọng không kém việc liệt kê chức năng.",
        ]),
        ("5. Chuẩn bị luận điểm vấn đáp", [
            "Khi vấn đáp về DFS, câu hỏi thường gặp là tại sao chọn DFS nếu thuật toán không tối ưu. Câu trả lời nên là DFS được chọn để đại diện cho nhóm tìm kiếm không thông tin, giúp minh họa cách duyệt sâu và quay lui trong không gian trạng thái. DFS không phải thuật toán tìm tuyến chính của hệ thống, nhưng là nền tảng để hiểu vì sao cần heuristic ở A*.",
            "Khi vấn đáp về A*, cần giải thích rõ điều kiện tối ưu. A* không tự động tối ưu trong mọi tình huống; nó tối ưu khi heuristic admissible và chi phí cạnh phù hợp. Nếu heuristic quá tham lam hoặc đánh giá quá cao, thuật toán có thể mất bảo đảm tối ưu. Câu trả lời tốt nên nhắc đến vai trò của g(n), h(n), open set, closed set và parent map.",
            "Khi vấn đáp về Simple Hill Climbing, cần tránh nói thuật toán luôn tìm lời giải tốt nhất. Nên thừa nhận thuật toán dễ kẹt local optimum, plateau hoặc ridge. Điểm mạnh của nó là đơn giản, ít tốn bộ nhớ và phù hợp để cải thiện lời giải ban đầu. Nếu giảng viên hỏi vì sao không chọn Simulated Annealing, có thể trả lời rằng yêu cầu mới chỉ chọn một thuật toán mỗi nhóm nên nhóm chọn phiên bản đơn giản để phân tích rõ bản chất local search.",
            "Khi vấn đáp về AND-OR Search, cần nhấn mạnh kết quả của thuật toán là conditional plan. Đây là điểm khác biệt lớn nhất so với A*. A* trả một đường đi trên graph đã biết; AND-OR Search trả kế hoạch có nhánh cho các outcome khác nhau. Nếu môi trường deterministic, AND-OR có thể không cần thiết; nếu môi trường bất định, nó biểu diễn được sự không chắc chắn tốt hơn path đơn.",
            "Khi vấn đáp về Forward Checking, cần nói rõ thuật toán thuộc nhóm CSP chứ không phải thuật toán tìm đường hình học. Nó làm việc với biến, domain và ràng buộc. Trong Find Your Path, Forward Checking phù hợp để loại các thứ tự giao hàng hoặc phân công chắc chắn vi phạm deadline, capacity hoặc pickup/dropoff. Vì vậy, thuật toán này bổ sung cho A*, không cạnh tranh trực tiếp với A*.",
            "Khi vấn đáp về Alpha-Beta, cần giải thích rằng Alpha-Beta giữ nguyên kết quả Minimax nhưng giảm số node cần duyệt. Alpha và beta là hai ngưỡng giúp nhận biết nhánh nào không còn khả năng ảnh hưởng đến quyết định cuối cùng. Trong bài toán Find Your Path, thuật toán được dùng như mô hình phân tích kịch bản bất lợi, không phải engine định tuyến chính.",
            "Một luận điểm tổng quát nên chuẩn bị là sáu thuật toán không được chọn ngẫu nhiên. Chúng tạo thành một bản đồ kiến thức của học phần: tìm kiếm mù, tìm kiếm có thông tin, tìm kiếm cục bộ, lập kế hoạch trong môi trường bất định, thỏa ràng buộc và tìm kiếm đối kháng. Nhờ vậy, báo cáo chứng minh được khả năng liên hệ giữa lý thuyết và bài toán ứng dụng.",
        ]),
        ("6. Rủi ro khi trình bày và cách kiểm soát", [
            "Rủi ro đầu tiên là dùng thuật ngữ không chính xác. Ví dụ, nếu gọi Online Re-planning là thuật toán độc lập trong khi thực chất nó là cơ chế tổng quát, giảng viên có thể bắt bẻ. Báo cáo hiện tại đã tránh rủi ro đó bằng cách chọn AND-OR Search cho nhóm môi trường phức tạp. Đây là thuật toán rõ ràng hơn và phù hợp với kiến thức thường gặp trong học phần AI.",
            "Rủi ro thứ hai là trình bày quá nhiều thuật toán nhưng phân tích mỏng. Khi mỗi nhóm chỉ chọn một thuật toán, báo cáo có điều kiện đi sâu vào trạng thái, mục tiêu, quy trình, công thức, độ phức tạp và tính chất. Cách này tốt hơn việc liệt kê hai thuật toán mỗi nhóm nhưng chỉ viết vài dòng cho mỗi thuật toán.",
            "Rủi ro thứ ba là phần đánh giá trông giống văn bản sinh tự động nếu chỉ dùng bảng. Bảng có thể tiện, nhưng dễ làm mất lập luận. Vì vậy, phần thảo luận đã chuyển sang văn xuôi học thuật, trong đó mỗi thuật toán được đặt vào vai trò cụ thể của hệ thống. Cách viết này tự nhiên hơn và phù hợp với báo cáo cuối kỳ hơn.",
            "Rủi ro thứ tư là hình minh họa không khớp nội dung. Nếu báo cáo nói về f(n), g(n), h(n) nhưng hình A* chỉ hiển thị đường đi cuối cùng, người đọc sẽ không thấy thuật toán hoạt động. Nếu báo cáo nói Forward Checking cắt domain nhưng hình không thể hiện domain, minh họa sẽ yếu. Vì vậy, mỗi hình cần được thiết kế theo đúng khái niệm thuật toán.",
            "Rủi ro thứ năm là số liệu thực nghiệm không ổn định. Thời gian chạy có thể thay đổi theo máy, dữ liệu và cách cài đặt. Để kiểm soát, nhóm nên dùng thêm chỉ tiêu độc lập hơn như số node mở rộng, số neighbor được xét, số giá trị bị loại khỏi domain hoặc số nhánh Alpha-Beta cắt được. Các chỉ tiêu này gắn chặt với bản chất thuật toán.",
            "Rủi ro cuối cùng là báo cáo quá thiên về mô tả project mà thiếu lý thuyết, hoặc quá thiên về lý thuyết mà thiếu liên hệ project. Bản báo cáo cần giữ cân bằng: mỗi thuật toán đều phải có định nghĩa và phân tích, nhưng cũng phải chỉ ra nó nằm ở đâu trong Find Your Path. Đây là điểm làm báo cáo có tính học thuật mà vẫn gắn với sản phẩm.",
        ]),
    ]

    for title, paragraphs in sections:
        add_heading_before(conclusion, title)
        for paragraph in paragraphs:
            add_para_before(conclusion, paragraph)

    remove_static_toc(doc)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
