import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DEFAULT_INPUT = Path("docs/find-your-path-report.docx")
DEFAULT_OUTPUT = Path("docs/find-your-path-report-reordered.docx")


def style_name(paragraph):
    return paragraph.style.name if paragraph.style is not None else ""


def heading_texts(doc):
    return [(i, style_name(p), p.text) for i, p in enumerate(doc.paragraphs) if style_name(p).startswith("Heading")]


def find_heading(doc, style, predicate, paragraphs=None):
    scope = paragraphs if paragraphs is not None else doc.paragraphs
    matches = [p for p in scope if style_name(p) == style and predicate(p.text)]
    if len(matches) != 1:
        details = "\n".join(f"{i}: {s} | {t}" for i, s, t in heading_texts(doc))
        raise RuntimeError(f"Expected one heading {style}, found {len(matches)}.\n{details}")
    return matches[0]


def paragraph_index(doc, paragraph):
    target = paragraph._p
    for i, p in enumerate(doc.paragraphs):
        if p._p is target:
            return i
    raise RuntimeError("Paragraph not found")


def collect_block(doc, start_heading, stop_style, stop_predicate):
    start = paragraph_index(doc, start_heading)
    end = len(doc.paragraphs)
    for i in range(start + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if style_name(p) == "Heading 2" or (style_name(p) == stop_style and stop_predicate(p.text)):
            end = i
            break
    return [p._p for p in doc.paragraphs[start:end]]


def scoped_paragraphs(doc, start_style, start_predicate, stop_style, stop_predicate):
    start = paragraph_index(doc, find_heading(doc, start_style, start_predicate))
    stop = paragraph_index(doc, find_heading(doc, stop_style, stop_predicate))
    if stop <= start:
        raise RuntimeError("Invalid scoped paragraph range")
    return doc.paragraphs[start:stop + 1]


def reorder_blocks(doc, scope, block_specs, stop_style, stop_predicate):
    blocks = []
    for predicate in block_specs:
        start = find_heading(doc, "Heading 2", predicate, scope)
        blocks.append(collect_block(doc, start, stop_style, stop_predicate))

    # Remove old blocks from the document body first.
    for block in blocks:
        for element in block:
            element.getparent().remove(element)

    anchor = find_heading(doc, stop_style, stop_predicate, scope)._p
    parent = anchor.getparent()
    insert_at = parent.index(anchor)
    for block in blocks:
        for element in block:
            parent.insert(insert_at, element)
            insert_at += 1


def replace_paragraph_text(paragraph, text):
    # Keep paragraph style/direct formatting while replacing visible text.
    paragraph.clear()
    paragraph.add_run(text)


def renumber_headings(doc):
    replacements = {
        "4. Nhóm thỏa ràng buộc: Backtracking và Forward Checking": "5. Nhóm thỏa ràng buộc: Backtracking và Forward Checking",
        "4.1. Backtracking": "5.1. Backtracking",
        "4.2. Forward Checking": "5.2. Forward Checking",
        "5. Nhóm tìm kiếm đối kháng: Minimax và Alpha-Beta Pruning": "6. Nhóm tìm kiếm đối kháng: Minimax và Alpha-Beta Pruning",
        "5.1. Minimax": "6.1. Minimax",
        "5.2. Alpha-Beta Pruning": "6.2. Alpha-Beta Pruning",
        "6. Nhóm môi trường phức tạp: Online Re-planning và AND-OR Search": "4. Nhóm môi trường phức tạp: Online Re-planning và AND-OR Search",
        "6.1. Online Re-planning": "4.1. Online Re-planning",
        "6.2. AND-OR Search": "4.2. AND-OR Search",
        "4. Nhóm thỏa ràng buộc": "5. Nhóm thỏa ràng buộc",
        "5. Nhóm tìm kiếm đối kháng": "6. Nhóm tìm kiếm đối kháng",
        "6. Nhóm môi trường phức tạp": "4. Nhóm môi trường phức tạp",
    }
    for paragraph in doc.paragraphs:
        if style_name(paragraph).startswith("Heading") and paragraph.text in replacements:
            replace_paragraph_text(paragraph, replacements[paragraph.text])


def update_discussion_order(doc):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Phần đánh giá không nên chỉ kết luận"):
            replace_paragraph_text(
                paragraph,
                "Phần đánh giá không nên chỉ kết luận thuật toán nào nhanh hơn, vì mỗi nhóm thuật toán trả lời một loại câu hỏi khác nhau. "
                "BFS và DFS giúp hiểu cách duyệt không gian trạng thái; Greedy và A* khai thác heuristic để giảm tìm kiếm mù; "
                "Hill Climbing và Simulated Annealing xử lý tối ưu cấu hình; Online Re-planning và AND-OR Search xử lý môi trường không quan sát đầy đủ; "
                "Backtracking và Forward Checking kiểm tra tính hợp lệ dưới ràng buộc; Minimax và Alpha-Beta phân tích quyết định trong tình huống có phản ứng bất lợi.",
            )
        elif paragraph.text.startswith("Quan điểm của nhóm là không nên ép"):
            replace_paragraph_text(
                paragraph,
                "Quan điểm của nhóm là không nên ép một thuật toán xử lý toàn bộ bài toán Find Your Path. Một thiết kế hợp lý có thể kết hợp nhiều lớp: "
                "thuật toán tìm đường xử lý graph, thuật toán cục bộ cải thiện thứ tự giao, thuật toán môi trường phức tạp thích nghi với quan sát thiếu, "
                "thuật toán ràng buộc kiểm tra tính khả thi của đơn hàng và mô hình đối kháng dùng để đánh giá kịch bản xấu. "
                "Cách kết hợp này phản ánh đúng tinh thần của trí tuệ nhân tạo: lựa chọn biểu diễn và thuật toán phụ thuộc vào bản chất của bài toán.",
            )
        elif paragraph.text.startswith("Báo cáo đã được tổ chức"):
            replace_paragraph_text(
                paragraph,
                "Báo cáo đã được tổ chức theo bố cục cuối kỳ gồm bài toán đặt ra, thuật toán áp dụng, thực nghiệm và kết quả, đánh giá và thảo luận, kết luận và tài liệu tham khảo. "
                "Phần bài toán được phân tích thông qua không gian trạng thái, dữ liệu vào ra, tiêu chí chi phí và mô hình PEAS. "
                "Phần thuật toán lựa chọn tối thiểu hai đại diện cho mỗi nhóm chính theo thứ tự: BFS/DFS, Greedy/A*, Hill Climbing/Simulated Annealing, "
                "Online Re-planning/AND-OR Search, Backtracking/Forward Checking và Minimax/Alpha-Beta Pruning.",
            )


def update_comparison_table(doc):
    for table in doc.tables:
        found = {}
        header = None
        for idx, row in enumerate(table.rows):
            first = row.cells[0].text.strip()
            if first == "Nhóm":
                header = row._tr
            elif "Không thông tin" in first:
                found["uninformed"] = row._tr
            elif "Có thông tin" in first:
                found["informed"] = row._tr
            elif "Cục bộ" in first:
                found["local"] = row._tr
            elif "Môi trường phức tạp" in first:
                found["complex"] = row._tr
            elif "Thỏa ràng buộc" in first or "Ràng buộc" in first:
                found["csp"] = row._tr
            elif "Đối kháng" in first:
                found["adversarial"] = row._tr
        ordered_keys = ["uninformed", "informed", "local", "complex", "csp", "adversarial"]
        if header is not None and set(ordered_keys).issubset(found):
            parent = header.getparent()
            ordered_rows = [found[key] for key in ordered_keys]
            for row in ordered_rows:
                parent.remove(row)
            for row in ordered_rows:
                parent.append(row)
            break


def mark_update_fields(doc):
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main():
    input_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_INPUT
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_OUTPUT
    if input_path.resolve() != output_path.resolve():
        shutil.copyfile(input_path, output_path)

    doc = Document(output_path)

    # Section II: Uninformed, Informed, Local, Complex, CSP, Adversarial.
    section_ii = scoped_paragraphs(
        doc,
        "Heading 1",
        lambda t: t.startswith("II."),
        "Heading 1",
        lambda t: t.startswith("III."),
    )
    reorder_blocks(
        doc,
        section_ii,
        [
            lambda t: "không thông tin" in t,
            lambda t: "có thông tin" in t,
            lambda t: "cục bộ" in t,
            lambda t: "môi trường phức tạp" in t,
            lambda t: "thỏa ràng buộc" in t,
            lambda t: "đối kháng" in t,
        ],
        "Heading 1",
        lambda t: t.startswith("III."),
    )

    # Section III: same group order; keep Link GitHub after all experiment groups.
    section_iii = scoped_paragraphs(
        doc,
        "Heading 1",
        lambda t: t.startswith("III."),
        "Heading 1",
        lambda t: t.startswith("IV."),
    )
    reorder_blocks(
        doc,
        section_iii,
        [
            lambda t: t.startswith("1.") and "không thông tin" in t,
            lambda t: t.startswith("2.") and "có thông tin" in t,
            lambda t: t.startswith("3.") and "cục bộ" in t,
            lambda t: "môi trường phức tạp" in t,
            lambda t: "thỏa ràng buộc" in t,
            lambda t: "đối kháng" in t,
            lambda t: "Link GitHub" in t,
        ],
        "Heading 1",
        lambda t: t.startswith("IV."),
    )

    renumber_headings(doc)
    update_discussion_order(doc)
    update_comparison_table(doc)
    mark_update_fields(doc)
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
