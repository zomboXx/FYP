from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path("docs/find-your-path-report.docx")


def style_name(paragraph):
    return paragraph.style.name if paragraph.style is not None else ""


def set_font(run, bold=None):
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    rpr.get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)
    if bold is not None:
        run.bold = bold


def format_para(paragraph, first_line=True):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(1)
    for run in paragraph.runs:
        set_font(run)


def add_para_before(anchor, text):
    p = anchor.insert_paragraph_before(text)
    format_para(p)
    return p


def add_heading_before(anchor, text, level):
    p = anchor.insert_paragraph_before("", style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def find_heading(doc, text):
    matches = [p for p in doc.paragraphs if style_name(p).startswith("Heading") and p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one heading {text!r}, found {len(matches)}")
    return matches[0]


def delete_section(doc, start_text):
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == start_text:
            start = i
            break
    if start is None:
        return
    for p in list(doc.paragraphs[start:]):
        p._p.getparent().remove(p._p)


def update_toc(doc):
    # Keep front matter clean; static page values are conservative and Word can update fields on open.
    toc_entries = [
        ("I. Bài toán đặt ra", "1"),
        ("1. Bài toán gì", "1"),
        ("2. Phân tích PEAS của bài toán", "3"),
        ("II. Thuật toán áp dụng", "6"),
        ("1. Nhóm tìm kiếm không thông tin: Depth-First Search (DFS)", "6"),
        ("2. Nhóm tìm kiếm có thông tin: A*", "10"),
        ("3. Nhóm tìm kiếm cục bộ: Simple Hill Climbing", "14"),
        ("4. Nhóm môi trường phức tạp: AND-OR Search", "18"),
        ("5. Nhóm thỏa ràng buộc: Forward Checking", "22"),
        ("6. Nhóm tìm kiếm đối kháng: Alpha-Beta Pruning", "26"),
        ("III. Thực nghiệm và kết quả", "30"),
        ("1. Thiết kế kịch bản thực nghiệm", "30"),
        ("2. Chỉ tiêu quan sát và cách ghi nhận kết quả", "32"),
        ("3. Định hướng chèn hình minh họa", "34"),
        ("4. Link GitHub", "35"),
        ("IV. Đánh giá và thảo luận", "36"),
        ("1. Tổng hợp vai trò của sáu thuật toán", "37"),
        ("2. Mối liên hệ giữa biểu diễn trạng thái và thuật toán", "39"),
        ("3. Hạn chế và hướng hoàn thiện thực nghiệm", "41"),
        ("V. Kết luận", "43"),
        ("Tài liệu tham khảo", "44"),
    ]
    toc_paras = [p for p in doc.paragraphs if style_name(p).startswith("toc")]
    for p, (label, page) in zip(toc_paras, toc_entries):
        p.text = f"{label}\t{page}"


def main():
    doc = Document(DOCX)
    delete_section(doc, "Phụ lục phân tích mở rộng")

    conclusion = find_heading(doc, "V. Kết luận")
    add_heading_before(conclusion, "1. Tổng hợp vai trò của sáu thuật toán", 2)
    for text in [
        "Sáu thuật toán trong báo cáo tạo thành một chuỗi tăng dần về mức độ biểu diễn tri thức. DFS bắt đầu từ giả định ít thông tin nhất: tác tử chỉ biết cấu trúc kề của trạng thái và mở rộng theo chiều sâu. A* bổ sung heuristic, làm cho quá trình tìm kiếm có định hướng hơn. Simple Hill Climbing chuyển trọng tâm từ tìm đường sang cải thiện cấu hình. AND-OR Search đưa vào khả năng hành động có nhiều outcome. Forward Checking nhấn mạnh ràng buộc, còn Alpha-Beta Pruning mô hình hóa phản ứng bất lợi.",
        "Nếu nhìn theo hệ thống, các thuật toán này không nên được xem như sáu lựa chọn độc lập để thay thế nhau. Chúng có thể nằm ở các tầng khác nhau của cùng một ứng dụng. Một tầng định tuyến dùng A* để tìm đường giữa hai điểm; một tầng kiểm tra dùng Forward Checking để loại phương án vi phạm; một tầng tối ưu cục bộ dùng Simple Hill Climbing để cải thiện thứ tự; một tầng phân tích rủi ro dùng AND-OR hoặc Alpha-Beta để khảo sát tình huống khó.",
        "Điều này giúp báo cáo tránh lỗi thường gặp là so sánh thuật toán theo một tiêu chí đơn giản như nhanh hay chậm. DFS có thể nhanh trên graph nhỏ nhưng không tối ưu. A* có thể tối ưu nhưng tốn bộ nhớ. Simple Hill Climbing có thể cải thiện nhanh nhưng dễ kẹt. AND-OR Search giàu mô hình nhưng dễ bùng nổ nhánh. Forward Checking cắt ràng buộc tốt nhưng không tự tối ưu chi phí. Alpha-Beta giảm duyệt cây nhưng cần mô hình utility hợp lý.",
        "Từ góc nhìn học thuật, điểm đáng giá không phải là chứng minh một thuật toán thắng toàn bộ các thuật toán còn lại, mà là xác định đúng hoàn cảnh sử dụng. Một thuật toán chỉ có ý nghĩa khi mô hình bài toán phù hợp với giả định của nó. A* cần heuristic; Forward Checking cần biến, miền giá trị và ràng buộc; AND-OR Search cần transition model có nhiều outcome; Alpha-Beta cần cây quyết định và hàm utility.",
    ]:
        add_para_before(conclusion, text)

    add_heading_before(conclusion, "2. Mối liên hệ giữa biểu diễn trạng thái và thuật toán", 2)
    for text in [
        "Biểu diễn trạng thái là cầu nối giữa bài toán thực tế và thuật toán. Nếu trạng thái chỉ gồm vị trí hiện tại, ta có bài toán path finding cổ điển. Nếu trạng thái gồm thêm tập đơn hàng chưa giao, thời gian và tải hiện tại, không gian trạng thái mở rộng rất nhanh. Nếu trạng thái còn chứa thông tin bất định hoặc belief về môi trường, bài toán trở thành lập kế hoạch trong môi trường phức tạp.",
        "DFS và A* dùng biểu diễn trạng thái tương đối gần với graph. Điểm khác biệt là DFS không có thước đo định hướng, còn A* thêm hàm h(n). Vì vậy, chỉ cần thay đổi biểu diễn chi phí và heuristic, hành vi của A* có thể thay đổi đáng kể. Một heuristic tốt giúp giảm node mở rộng; một heuristic sai có thể làm mất bảo đảm tối ưu.",
        "Simple Hill Climbing lại yêu cầu trạng thái ở mức cấu hình. Thay vì hỏi node kế tiếp là gì, thuật toán hỏi route hiện tại có thể cải thiện bằng neighbor nào. Do đó, việc định nghĩa neighbor quan trọng không kém hàm cost. Nếu neighbor quá nghèo, thuật toán dễ kẹt; nếu neighbor quá rộng, chi phí mỗi vòng tăng cao.",
        "Forward Checking yêu cầu biểu diễn biến và domain. Đây là cách nhìn khác hẳn graph search. Một đơn hàng không còn chỉ là điểm trên bản đồ, mà trở thành biến cần được gán vị trí, thời điểm hoặc shipper. Ràng buộc chuyển thành điều kiện loại giá trị khỏi domain. Nhờ cách biểu diễn này, nhiều phương án sai bị loại trước khi cần tính route chi tiết.",
        "AND-OR Search và Alpha-Beta Pruning đòi hỏi mô hình hóa nhánh quyết định. AND-OR cần biết hành động nào có thể tạo outcome nào. Alpha-Beta cần xác định lượt MAX, lượt MIN và utility ở lá. Nếu không xây dựng được mô hình trạng thái phù hợp, hai thuật toán này sẽ trở thành mô tả hình thức. Vì vậy, trong demo, nhóm cần chọn kịch bản nhỏ nhưng rõ để minh họa đúng bản chất.",
    ]:
        add_para_before(conclusion, text)

    add_heading_before(conclusion, "3. Hạn chế và hướng hoàn thiện thực nghiệm", 2)
    for text in [
        "Hạn chế lớn nhất của báo cáo hiện tại là phần thực nghiệm phụ thuộc vào mức độ hoàn thiện của chương trình và dữ liệu demo. Phân tích thuật toán đã đủ để trình bày lý thuyết, nhưng để thuyết phục trong báo cáo cuối kỳ, nhóm cần bổ sung hình ảnh hoặc animation cho từng thuật toán. Mỗi hình nên thể hiện quá trình chứ không chỉ kết quả cuối cùng.",
        "Đối với DFS, hướng hoàn thiện là xuất trace cho từng bước mở rộng, bao gồm stack và visited set. Đối với A*, cần xuất open set, closed set, g-score và f-score. Đối với Simple Hill Climbing, cần xuất cost qua từng vòng. Đối với AND-OR Search, cần vẽ conditional plan. Đối với Forward Checking, cần vẽ domain bị thu hẹp. Đối với Alpha-Beta, cần đánh dấu nhánh bị cắt.",
        "Một hướng hoàn thiện khác là thống nhất dữ liệu đầu vào giữa các thuật toán ở mức có thể. DFS và A* có thể dùng cùng graph để so sánh tìm kiếm mù và tìm kiếm có thông tin. Simple Hill Climbing có thể dùng output route từ A* làm một phần của cấu hình ban đầu. Forward Checking có thể kiểm tra ràng buộc trên cùng tập đơn hàng. Cách liên kết này giúp báo cáo giống một dự án tổng hợp thay vì sáu ví dụ rời rạc.",
        "Trong phần vấn đáp, nhóm nên chuẩn bị câu trả lời về lý do không chọn hai thuật toán cho mỗi nhóm. Lý do là yêu cầu mới của phần II tập trung vào mỗi mục một thuật toán, nên báo cáo chọn đại diện tiêu biểu và phân tích sâu hơn. Cách này giúp tránh tình trạng liệt kê nhiều thuật toán nhưng không đủ độ sâu về trạng thái, mục tiêu, độ phức tạp và tính chất.",
        "Cuối cùng, cần lưu ý rằng độ dài báo cáo không nên đến từ việc lặp lại định nghĩa, mà từ việc phân tích mối liên hệ giữa thuật toán và bài toán. Các đoạn mở rộng trong phần thảo luận phải giải thích được tại sao thuật toán phù hợp, giới hạn của nó là gì và nếu đưa vào hệ thống thực tế thì nên đặt ở tầng nào. Đây là cách làm cho báo cáo dài hơn nhưng vẫn có giá trị học thuật.",
    ]:
        add_para_before(conclusion, text)

    update_toc(doc)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
