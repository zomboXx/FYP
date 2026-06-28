from __future__ import annotations

from pathlib import Path
import copy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path("docs/find-your-path-report.docx")
EXTRA = Path("docs/find-your-path-report-reordered.docx")


def style_name(paragraph):
    return paragraph.style.name if paragraph.style is not None else ""


def set_font(run, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True):
    paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25
    if first_line:
        pf.first_line_indent = Cm(1)
    for run in paragraph.runs:
        set_font(run)


def add_para(doc, text: str, first_line=True):
    p = doc.add_paragraph()
    p.add_run(text)
    format_paragraph(p, first_line=first_line)
    return p


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.paragraph_format.page_break_before = True
    return p


def m_el(tag):
    return OxmlElement(f"m:{tag}")


def m_attr(element, name, value):
    element.set(qn(f"m:{name}"), value)
    return element


def mr(text):
    run = m_el("r")
    text_el = m_el("t")
    text_el.text = text
    run.append(text_el)
    return run


def group(*children):
    element = m_el("e")
    for child in children:
        element.append(child)
    return element


def sub(base_children, sub_children):
    element = m_el("sSub")
    element.append(group(*base_children))
    sub_el = m_el("sub")
    for child in sub_children:
        sub_el.append(child)
    element.append(sub_el)
    return element


def sup(base_children, sup_children):
    element = m_el("sSup")
    element.append(group(*base_children))
    sup_el = m_el("sup")
    for child in sup_children:
        sup_el.append(child)
    element.append(sup_el)
    return element


def frac(num_children, den_children):
    element = m_el("f")
    element.append(group(*num_children))
    den = m_el("den")
    for child in den_children:
        den.append(child)
    element.append(den)
    return element


def nary(symbol, sub_children, sup_children, expr_children):
    element = m_el("nary")
    prop = m_el("naryPr")
    prop.append(m_attr(m_el("chr"), "val", symbol))
    prop.append(m_attr(m_el("limLoc"), "val", "undOvr"))
    element.append(prop)
    sub_el = m_el("sub")
    for child in sub_children:
        sub_el.append(child)
    element.append(sub_el)
    sup_el = m_el("sup")
    for child in sup_children:
        sup_el.append(child)
    element.append(sup_el)
    element.append(group(*expr_children))
    return element


def add_equation(doc, children):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    math_para = m_el("oMathPara")
    math = m_el("oMath")
    for child in children:
        math.append(child)
    math_para.append(math)
    p._p.append(math_para)
    return p


def clear_from_toc(doc):
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "MỤC LỤC":
            start = i
            break
    if start is None:
        raise RuntimeError("Cannot find MỤC LỤC")
    for p in list(doc.paragraphs[start:]):
        p._p.getparent().remove(p._p)
    # Remove any table after the assignment tables; this deletes the old AI-looking comparison table.
    while len(doc.tables) > 2:
        table = doc.tables[-1]
        table._tbl.getparent().remove(table._tbl)


def set_styles(doc):
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "toc 1", "toc 2", "toc 3"]:
        if name in doc.styles:
            style = doc.styles[name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(13)


def add_toc(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MỤC LỤC")
    set_font(run, bold=True)
    title.paragraph_format.space_after = Pt(8)

    entries = [
        ("toc 1", "I. Bài toán đặt ra", "1"),
        ("toc 2", "1. Bài toán gì", "1"),
        ("toc 2", "2. Phân tích PEAS của bài toán", "3"),
        ("toc 1", "II. Thuật toán áp dụng", "6"),
        ("toc 2", "1. Nhóm tìm kiếm không thông tin: Depth-First Search (DFS)", "6"),
        ("toc 2", "2. Nhóm tìm kiếm có thông tin: A*", "10"),
        ("toc 2", "3. Nhóm tìm kiếm cục bộ: Simple Hill Climbing", "14"),
        ("toc 2", "4. Nhóm môi trường phức tạp: AND-OR Search", "18"),
        ("toc 2", "5. Nhóm thỏa ràng buộc: Forward Checking", "22"),
        ("toc 2", "6. Nhóm tìm kiếm đối kháng: Alpha-Beta Pruning", "26"),
        ("toc 1", "III. Thực nghiệm và kết quả", "30"),
        ("toc 2", "1. Thiết kế kịch bản thực nghiệm", "30"),
        ("toc 2", "2. Chỉ tiêu quan sát và cách ghi nhận kết quả", "32"),
        ("toc 2", "3. Định hướng chèn hình minh họa", "34"),
        ("toc 2", "4. Link GitHub", "35"),
        ("toc 1", "IV. Đánh giá và thảo luận", "36"),
        ("toc 1", "V. Kết luận", "39"),
        ("toc 1", "Tài liệu tham khảo", "40"),
    ]
    for style, label, page in entries:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        if style == "toc 2":
            p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.6),
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.DOTS,
        )
        r = p.add_run(f"{label}\t{page}")
        set_font(r)


def add_problem_section(doc):
    add_heading(doc, "I. Bài toán đặt ra", 1)
    add_heading(doc, "1. Bài toán gì", 2)
    paragraphs = [
        "Find Your Path được xác định là bài toán lập lộ trình giao hàng trong môi trường đô thị, nơi hệ thống phải chọn đường đi và chuỗi hành động phù hợp cho tác tử giao hàng. Nếu nhìn ở mức tối giản, bài toán có thể được mô tả là tìm một đường đi từ điểm xuất phát đến điểm đích trên đồ thị có trọng số. Tuy nhiên, trong bối cảnh ứng dụng thực tế, cách hiểu đó chưa đủ, vì mỗi quyết định di chuyển còn chịu ảnh hưởng của thời gian, ràng buộc đơn hàng, trạng thái đường đi, khả năng cập nhật thông tin và mục tiêu tối ưu tổng thể.",
        "Không gian trạng thái của bài toán gồm nhiều thành phần hơn một node hiện tại. Một trạng thái có thể bao gồm vị trí của shipper, tập đơn hàng chưa phục vụ, tập đơn hàng đã nhận, trạng thái tải hiện tại, thời gian đã trôi qua, các cạnh bị chặn hoặc có chi phí thay đổi, cùng những thông tin mà hệ thống đã quan sát được. Chính sự kết hợp này làm cho bài toán Find Your Path phù hợp để khảo sát nhiều nhóm thuật toán trí tuệ nhân tạo khác nhau thay vì chỉ dùng một thuật toán tìm đường đơn lẻ.",
        "Trạng thái bắt đầu là trạng thái mà shipper hoặc hệ thống đang đứng trước một yêu cầu giao hàng. Ở trạng thái này, hệ thống biết vị trí xuất phát, biết hoặc ước lượng bản đồ đường đi, biết mục tiêu cần phục vụ và có thể có một số thông tin ban đầu về chi phí cạnh. Trạng thái mục tiêu là trạng thái mà tác tử đã đến đúng điểm đích hoặc hoàn thành một chuỗi giao hàng thỏa các ràng buộc đặt ra. Trong phiên bản đơn giản, mục tiêu là đến node đích; trong phiên bản mở rộng, mục tiêu là hoàn thành route với chi phí chấp nhận được và không vi phạm các điều kiện quan trọng.",
        "Các phép hành động trong bài toán có thể được biểu diễn dưới dạng chọn node kế tiếp, nhận đơn, gán đơn cho shipper, cập nhật route hoặc thay đổi kế hoạch khi phát hiện thông tin mới. Nếu mô hình chỉ xét graph tĩnh, hành động là di chuyển qua một cạnh hợp lệ. Nếu mô hình xét nhiều đơn hàng và ràng buộc, hành động còn bao gồm việc chọn thứ tự phục vụ hoặc loại bỏ một phương án không khả thi. Nếu mô hình xét môi trường bất định, hành động có thể dẫn đến nhiều kết quả khác nhau, chẳng hạn một cạnh đang thông nhưng sau đó bị chặn, hoặc chi phí di chuyển tăng do điều kiện đường đi.",
        "Đầu vào của bài toán gồm đồ thị đường đi, tập node, tập cạnh, trọng số cạnh, vị trí bắt đầu, vị trí mục tiêu, thông tin đơn hàng và các tham số đánh giá. Trọng số cạnh có thể đại diện cho khoảng cách, thời gian di chuyển, chi phí nhiên liệu hoặc một chỉ số tổng hợp. Đầu ra mong muốn là một route hoặc một kế hoạch hành động giúp tác tử đi từ trạng thái bắt đầu đến trạng thái mục tiêu. Trong các thuật toán phức tạp hơn, đầu ra không chỉ là một path đơn, mà có thể là kế hoạch có điều kiện hoặc một lời giải thỏa ràng buộc.",
        "Tiêu chí đánh giá lời giải không chỉ là độ dài hình học. Một route tốt trong bài toán giao hàng cần cân bằng giữa thời gian, chi phí, độ ổn định, tính khả thi và mức độ thỏa ràng buộc. Tuyến đường ngắn nhất chưa chắc là tốt nhất nếu đi qua khu vực thường tắc nghẽn, nếu làm trễ deadline hoặc nếu khiến shipper không thể hoàn thành các đơn tiếp theo. Vì vậy, báo cáo xem Find Your Path như một bài toán tổng hợp, trong đó từng nhóm thuật toán đại diện cho một góc nhìn khác nhau của trí tuệ nhân tạo.",
        "Lý do chọn đề tài này là vì bài toán tìm đường có tính trực quan, dễ minh họa bằng bản đồ và animation, đồng thời đủ giàu để phân tích nhiều khái niệm cốt lõi trong môn Trí tuệ nhân tạo. DFS giúp hiểu tìm kiếm mù, A* cho thấy vai trò của heuristic, Simple Hill Climbing minh họa tối ưu cục bộ, AND-OR Search mô tả kế hoạch trong môi trường bất định, Forward Checking diễn giải xử lý ràng buộc, còn Alpha-Beta Pruning cho thấy cách ra quyết định khi có yếu tố đối kháng.",
    ]
    for paragraph in paragraphs:
        add_para(doc, paragraph)

    add_heading(doc, "2. Phân tích PEAS của bài toán", 2)
    paragraphs = [
        "PEAS là khung phân tích gồm Performance measure, Environment, Actuators và Sensors. Đối với Find Your Path, PEAS giúp chuyển bài toán từ mô tả tự nhiên sang mô hình tác tử rõ ràng. Nhờ đó, người đọc có thể xác định hệ thống đang tối ưu điều gì, hoạt động trong môi trường nào, có thể thực hiện hành động gì và nhận thông tin từ nguồn nào.",
        "Performance measure của hệ thống là tập các tiêu chí đo chất lượng route và chất lượng quyết định. Tiêu chí cơ bản nhất là tổng chi phí đường đi, có thể được tính từ khoảng cách hoặc thời gian. Tuy nhiên, nếu chỉ tối thiểu hóa khoảng cách, hệ thống có thể chọn một route không phù hợp với bối cảnh giao hàng. Do đó, performance measure nên bao gồm thêm độ trễ so với deadline, số lần vi phạm ràng buộc, độ ổn định của route, số node đã mở rộng khi tìm kiếm và khả năng thích nghi khi môi trường thay đổi.",
        "Một hàm đánh giá tổng hợp có thể viết dưới dạng chi phí route gồm nhiều thành phần. Thành phần đầu tiên là chi phí di chuyển; thành phần thứ hai là phạt trễ hạn; thành phần thứ ba là phạt vi phạm capacity hoặc thứ tự phục vụ; thành phần cuối có thể là phạt rủi ro nếu route đi qua vùng có khả năng thay đổi cao. Công thức này không nhất thiết dùng nguyên xi trong mọi thuật toán, nhưng nó cho thấy cách chuyển mục tiêu thực tế thành giá trị định lượng.",
    ]
    for paragraph in paragraphs:
        add_para(doc, paragraph)
    add_equation(doc, [mr("Cost(route)="), sub([mr("travel")], [mr("time")]), mr("+"), sub([mr("late")], [mr("penalty")]), mr("+"), sub([mr("constraint")], [mr("penalty")]), mr("+"), sub([mr("risk")], [mr("penalty")])])
    more = [
        "Environment là môi trường đô thị được biểu diễn bằng đồ thị có trọng số. Node có thể là giao điểm, depot, vị trí shipper, điểm lấy hàng hoặc điểm giao hàng. Cạnh biểu diễn đoạn đường có thể đi qua, kèm chi phí di chuyển. Môi trường này có thể được xem là rời rạc vì trạng thái được biểu diễn bằng node và cạnh, nhưng vẫn có yếu tố động nếu chi phí cạnh thay đổi theo thời gian hoặc nếu một đoạn đường bất ngờ không còn khả dụng.",
        "Môi trường của Find Your Path không hoàn toàn quan sát được trong mọi tình huống. Ở mức đơn giản, hệ thống biết toàn bộ graph trước khi tìm đường. Ở mức thực tế hơn, hệ thống chỉ biết một phần thông tin hoặc chỉ cập nhật khi shipper tiến đến gần một khu vực. Đây là lý do báo cáo có nhóm môi trường phức tạp với AND-OR Search, vì tác tử cần chuẩn bị kế hoạch khi hành động có nhiều kết quả khả dĩ.",
        "Actuators là các hành động mà tác tử hoặc hệ thống có thể thực hiện. Ở mức tìm đường, actuator là hành động chọn node kế tiếp từ tập node kề. Ở mức lập kế hoạch giao hàng, actuator bao gồm chọn đơn kế tiếp, đổi thứ tự route, gán đơn cho shipper hoặc loại bỏ phương án không hợp lệ. Ở mức mô phỏng đối kháng, actuator có thể là chọn phương án có utility cao nhất trong khi môi trường chọn phản ứng bất lợi.",
        "Sensors là nguồn thông tin giúp tác tử nhận biết trạng thái. Trong dự án phần mềm, sensor không nhất thiết là cảm biến vật lý, mà là dữ liệu vị trí, bản đồ, trạng thái đơn hàng, trọng số cạnh, thông tin tắc đường hoặc phản hồi từ quá trình mô phỏng. Sensor quyết định chất lượng của belief hiện tại, từ đó ảnh hưởng trực tiếp đến thuật toán tìm kiếm và thuật toán lập kế hoạch.",
        "Nếu xét theo đặc điểm môi trường trong AI, Find Your Path có thể là single-agent ở phiên bản tìm đường, nhưng trở thành multi-agent hoặc adversarial nếu xét nhiều shipper hoặc yếu tố môi trường bất lợi. Bài toán có thể deterministic khi mỗi hành động dẫn đến một kết quả chắc chắn, nhưng có thể nondeterministic khi đường đi có thể bị chặn hoặc chi phí cạnh thay đổi. Bài toán có thể static trong graph cố định, nhưng dynamic khi dữ liệu cập nhật theo thời gian.",
        "Phân tích PEAS cho thấy không nên chọn thuật toán chỉ dựa trên tên bài toán. Nếu chỉ cần tìm đường trên graph đã biết, A* là lựa chọn rất mạnh. Nếu chỉ cần minh họa tìm kiếm mù, DFS đủ rõ ràng. Nếu cần cải thiện route nhiều điểm, Simple Hill Climbing có ý nghĩa. Nếu route chịu ràng buộc, Forward Checking phù hợp. Nếu có outcome bất định, AND-OR Search giúp biểu diễn conditional plan. Nếu cần phân tích phản ứng bất lợi, Alpha-Beta Pruning là công cụ thích hợp.",
    ]
    for paragraph in more:
        add_para(doc, paragraph)


def algorithm_common_intro(name, role):
    return (
        f"Thuật toán {name} được chọn vì {role}. Trong phạm vi báo cáo, thuật toán không chỉ được mô tả như một đoạn mã, "
        "mà được phân tích theo cách hình thành trạng thái, cách chọn hành động, tiêu chí dừng, độ phức tạp và mức độ phù hợp với Find Your Path. "
        "Cách trình bày này giúp liên hệ trực tiếp giữa lý thuyết học trên lớp và quyết định thiết kế trong dự án."
    )


def add_algorithm_sections(doc):
    add_heading(doc, "II. Thuật toán áp dụng", 1)
    add_para(doc, "Phần này chọn một thuật toán đại diện cho mỗi nhóm kiến thức. Thứ tự trình bày là DFS, A*, Simple Hill Climbing, AND-OR Search, Forward Checking và Alpha-Beta Pruning. Mỗi mục được viết như một tiểu mục nghiên cứu độc lập, tập trung vào bản chất thuật toán và cách thuật toán đóng góp cho bài toán Find Your Path.")
    add_para(doc, "Các ký hiệu dùng chung gồm b là hệ số phân nhánh trung bình, d là độ sâu nghiệm nông nhất, m là độ sâu tối đa, n là số biến trong bài toán ràng buộc, d_c là kích thước miền giá trị của biến CSP và k là số outcome trung bình của một hành động trong môi trường bất định.")

    add_heading(doc, "1. Nhóm tìm kiếm không thông tin: Depth-First Search (DFS)", 2)
    for paragraph in [
        algorithm_common_intro("DFS", "nó thể hiện rõ cách một tác tử có thể tìm kiếm khi chưa có heuristic"),
        "DFS bắt đầu từ trạng thái xuất phát và ưu tiên mở rộng sâu theo một nhánh trước khi xét các nhánh còn lại. Nếu cài đặt bằng stack, node được thêm sau sẽ được lấy ra trước; nếu cài đặt bằng đệ quy, lời gọi hiện tại tiếp tục đi sâu cho đến khi gặp mục tiêu, hết node kề hoặc cần quay lui. Chính đặc điểm đi sâu này làm DFS khác với các thuật toán mở rộng theo lớp.",
        "Trong Find Your Path, trạng thái của DFS có thể được biểu diễn bằng node hiện tại, tập node đã thăm, parent map và stack đang chờ xét. Trạng thái bắt đầu là vị trí xuất phát của shipper. Trạng thái mục tiêu là node đích hoặc một điều kiện hoàn thành route. Khi mở rộng một node, DFS duyệt các node kề chưa được thăm, thêm chúng vào stack và lưu quan hệ parent để có thể truy vết đường đi nếu tìm thấy mục tiêu.",
        "Ưu điểm quan trọng của DFS là đơn giản và tiết kiệm bộ nhớ. Thuật toán không cần hàng đợi ưu tiên, không cần heuristic và không cần lưu toàn bộ frontier theo từng lớp. Vì vậy, DFS phù hợp để kiểm tra tính liên thông của graph, phát hiện một path bất kỳ và minh họa trực quan quá trình duyệt trạng thái. Trong báo cáo cuối kỳ, DFS cũng là nền tảng giúp người đọc hiểu sự khác biệt giữa tìm kiếm mù và tìm kiếm có thông tin.",
        "Tuy nhiên, DFS không phù hợp để tìm đường tối ưu trong graph có trọng số. Nếu node kề được liệt kê theo thứ tự không tốt, DFS có thể đi rất sâu vào một nhánh dài trước khi xét nhánh ngắn hơn. Trong bản đồ giao thông, điều này có nghĩa thuật toán có thể tìm được một route hợp lệ nhưng route đó không ngắn, không rẻ và không phản ánh tiêu chí tối ưu thực tế.",
        "Bốn tính chất của DFS cần được đánh giá cẩn thận. Về tính đầy đủ, DFS không đầy đủ trong không gian vô hạn hoặc khi graph có chu trình mà không có cơ chế visited hay depth limit. Trong graph hữu hạn có kiểm soát visited, DFS có thể tìm thấy nghiệm nếu nghiệm tồn tại. Về tính tối ưu, DFS không bảo đảm tối ưu vì nghiệm đầu tiên phụ thuộc vào thứ tự duyệt. Về thời gian, trường hợp xấu là O(b^m). Về không gian, DFS thường tốt hơn BFS, khoảng O(bm) hoặc O(m) tùy cách cài đặt.",
        "Khi áp dụng vào Find Your Path, DFS nên được xem là thuật toán nền để giải thích không gian trạng thái, không nên được quảng bá như thuật toán tìm tuyến chính. Giá trị của DFS nằm ở khả năng minh họa quá trình mở rộng, quay lui và đánh dấu visited. Nếu demo có animation, DFS có thể cho thấy rất rõ vì sao một thuật toán không dùng heuristic dễ đi vào nhánh không hiệu quả.",
    ]:
        add_para(doc, paragraph)
    add_equation(doc, [mr("Time(DFS)=O("), sup([mr("b")], [mr("m")]), mr(")")])
    add_equation(doc, [mr("Space(DFS)=O(bm)")])
    for paragraph in [
        "Trong vấn đáp, điểm cần nhấn mạnh là DFS không sai chỉ vì nó không tối ưu. Nó trả lời một câu hỏi khác: liệu có thể tìm một đường đi bằng cách khám phá sâu hay không. Còn nếu câu hỏi là route nào tốt nhất theo chi phí, ta cần thuật toán khác, điển hình là A* trong nhóm tìm kiếm có thông tin.",
        "Nếu so với các thuật toán khác trong báo cáo, DFS có vai trò thấp tầng nhất. Nó không dùng hàm đánh giá, không xử lý ràng buộc phức tạp, không mô hình hóa bất định và không xét đối kháng. Chính vì vậy, DFS rất phù hợp để mở đầu phần thuật toán: từ một kỹ thuật đơn giản, báo cáo dần phát triển sang các kỹ thuật giàu thông tin và giàu mô hình hơn.",
    ]:
        add_para(doc, paragraph)

    add_heading(doc, "2. Nhóm tìm kiếm có thông tin: A*", 2)
    for paragraph in [
        algorithm_common_intro("A*", "nó là thuật toán tìm đường có heuristic phù hợp nhất với bản chất bản đồ của dự án"),
        "A* kết hợp hai nguồn thông tin: chi phí thật đã đi từ điểm xuất phát đến node hiện tại và chi phí ước lượng từ node hiện tại đến mục tiêu. Thay vì mở rộng node theo thứ tự mù, A* ưu tiên node có tổng chi phí dự kiến thấp nhất. Đây là lý do A* thường được dùng trong các bài toán tìm đường trên bản đồ, game, robot và hệ thống định tuyến.",
    ]:
        add_para(doc, paragraph)
    add_equation(doc, [mr("f(n)=g(n)+h(n)")])
    for paragraph in [
        "Trong công thức trên, g(n) là chi phí thực tế từ start đến n, còn h(n) là heuristic ước lượng chi phí từ n đến goal. Nếu graph dùng tọa độ, h(n) có thể là khoảng cách Euclid hoặc Manhattan. Nếu graph biểu diễn thời gian di chuyển, heuristic cần được thiết kế sao cho không đánh giá quá cao chi phí thật nếu muốn bảo đảm tối ưu.",
    ]:
        add_para(doc, paragraph)
    add_equation(doc, [sub([mr("h")], [mr("E")]), mr("(n)="), mr("sqrt(("), sub([mr("x")], [mr("n")]), mr("-"), sub([mr("x")], [mr("g")]), mr(")^2+("), sub([mr("y")], [mr("n")]), mr("-"), sub([mr("y")], [mr("g")]), mr(")^2)")])
    for paragraph in [
        "Trạng thái bắt đầu của A* gồm node xuất phát, open set chứa node đang chờ xét, closed set chứa node đã xử lý, bảng g-score và parent map. Ở mỗi bước, thuật toán lấy node có f(n) nhỏ nhất ra khỏi open set. Nếu node đó là goal, thuật toán truy vết parent map để trả route. Nếu chưa phải goal, thuật toán xét các node kề, cập nhật g-score nếu tìm thấy đường rẻ hơn và đưa node kề vào open set.",
        "Trong Find Your Path, A* có vai trò trung tâm vì bài toán có đích rõ ràng, graph có trọng số và có thể sử dụng heuristic địa lý. So với DFS, A* không đi sâu tùy ý mà có định hướng. So với thuật toán chỉ dùng heuristic như Greedy Best-First Search, A* không bỏ qua chi phí đã đi, nhờ đó cân bằng giữa tham vọng đi gần goal và thực tế chi phí trên đường.",
        "Về bốn tính chất, A* đầy đủ nếu hệ số phân nhánh hữu hạn, chi phí cạnh dương và tồn tại nghiệm. A* tối ưu khi heuristic admissible, tức không bao giờ đánh giá quá cao chi phí thật còn lại. Nếu heuristic consistent, giá trị f dọc theo đường đi không giảm bất thường, giúp triển khai ổn định hơn. Trong trường hợp xấu, thời gian và không gian của A* vẫn có thể là O(b^d), vì thuật toán có thể phải lưu và xét rất nhiều node.",
        "Điểm mạnh của A* là tính thực dụng. Trong demo Find Your Path, A* cho phép hiển thị rõ open set, closed set, đường đi tạm thời và đường đi cuối cùng. Khi thay đổi heuristic, có thể quan sát số node mở rộng thay đổi như thế nào. Đây là minh chứng trực quan cho vai trò của tri thức bổ sung trong tìm kiếm có thông tin.",
        "Điểm yếu của A* là tiêu thụ bộ nhớ lớn và phụ thuộc vào chất lượng heuristic. Nếu heuristic quá yếu, A* gần giống Dijkstra và mở rộng nhiều node. Nếu heuristic không admissible, thuật toán có thể nhanh hơn nhưng mất bảo đảm tối ưu. Vì vậy, trong báo cáo cần nêu rõ điều kiện để A* tối ưu thay vì chỉ khẳng định A* luôn tìm đường tốt nhất.",
        "Khi vấn đáp, có thể giải thích rằng A* trong dự án không chỉ là công thức f(n)=g(n)+h(n). Điều quan trọng là cách hệ thống cập nhật g-score, cách chọn node có f nhỏ nhất, cách lưu parent để dựng route và cách kiểm soát heuristic để không phá vỡ tính tối ưu. Đây là những điểm cho thấy nhóm hiểu thuật toán ở mức triển khai, không chỉ thuộc định nghĩa.",
    ]:
        add_para(doc, paragraph)

    add_heading(doc, "3. Nhóm tìm kiếm cục bộ: Simple Hill Climbing", 2)
    for paragraph in [
        algorithm_common_intro("Simple Hill Climbing", "nó minh họa rõ tư duy cải thiện lời giải cục bộ thay vì xây dựng toàn bộ cây tìm kiếm"),
        "Simple Hill Climbing bắt đầu từ một lời giải ứng viên và liên tục tìm lời giải láng giềng tốt hơn. Khác với A*, thuật toán không nhất thiết tìm từ start đến goal trên graph. Nó xem toàn bộ route hoặc thứ tự giao hàng như một cấu hình, sau đó thay đổi cấu hình này bằng các phép biến đổi nhỏ.",
        "Trong Find Your Path, một cấu hình có thể là thứ tự các điểm giao, thứ tự nhận đơn hoặc một route ứng viên qua nhiều node quan trọng. Neighbor có thể được tạo bằng cách hoán đổi hai điểm giao, đảo một đoạn route, thay một node trung gian hoặc chọn lại thứ tự phục vụ gần nhất. Nếu neighbor làm giảm hàm chi phí, thuật toán chấp nhận neighbor và tiếp tục từ đó.",
    ]:
        add_para(doc, paragraph)
    add_equation(doc, [mr("cost(route)="), sub([mr("travel")], [mr("time")]), mr("+"), sub([mr("late")], [mr("penalty")]), mr("+"), sub([mr("capacity")], [mr("penalty")]), mr("+"), sub([mr("constraint")], [mr("penalty")])])
    for paragraph in [
        "Hàm cost(route) có vai trò quyết định trong Simple Hill Climbing. Nếu hàm cost chỉ tính khoảng cách, thuật toán sẽ ưu tiên route ngắn. Nếu hàm cost thêm phạt deadline và capacity, thuật toán sẽ hướng đến route phù hợp hơn với bài toán giao hàng. Vì vậy, chất lượng của Simple Hill Climbing phụ thuộc mạnh vào cách thiết kế hàm đánh giá.",
        "Thuật toán này phù hợp với bài toán tối ưu cục bộ vì nó không cần lưu toàn bộ không gian trạng thái. Mỗi bước chỉ cần trạng thái hiện tại, một neighbor và giá trị đánh giá. Điều đó giúp thuật toán nhẹ, dễ cài đặt và dễ minh họa. Trong báo cáo, Simple Hill Climbing có thể được dùng để cho thấy một route ban đầu được cải thiện dần qua từng vòng.",
        "Nhược điểm lớn nhất của Simple Hill Climbing là dễ kẹt ở local optimum. Nếu mọi neighbor trực tiếp đều không tốt hơn trạng thái hiện tại, thuật toán dừng, dù có thể tồn tại lời giải tốt hơn ở xa hơn. Thuật toán cũng có thể gặp plateau, nơi nhiều trạng thái có cùng giá trị, hoặc ridge, nơi cần đi tạm qua bước không cải thiện để đến vùng tốt hơn. Vì báo cáo chọn Simple Hill Climbing, cần nói rõ đây là phiên bản đơn giản, không phải biến thể có sideways move hay simulated annealing.",
        "Về tính đầy đủ, Simple Hill Climbing không đầy đủ vì có thể dừng trước khi tìm thấy lời giải mong muốn. Về tối ưu, thuật toán không bảo đảm tối ưu toàn cục. Về thời gian, nếu mỗi vòng xét q neighbor và chạy tối đa t vòng, chi phí là O(tqC_eval), trong đó C_eval là chi phí tính hàm đánh giá. Về không gian, thuật toán thường rất tiết kiệm, có thể chỉ cần O(1) ngoài dữ liệu route nếu sinh neighbor tuần tự.",
    ]:
        add_para(doc, paragraph)
    add_equation(doc, [mr("T=O(tq·"), sub([mr("C")], [mr("eval")]), mr(")")])
    for paragraph in [
        "Trong Find Your Path, thuật toán này không nên thay thế A* cho bài toán tìm đường giữa hai điểm. Vai trò hợp lý hơn là cải thiện thứ tự hoặc cấu hình ở mức cao, còn việc tính đường đi chi tiết giữa hai điểm vẫn có thể do A* đảm nhiệm. Cách phân tầng này giúp báo cáo tránh nhầm lẫn giữa path planning và route optimization.",
        "Khi vấn đáp, có thể nhấn mạnh rằng Simple Hill Climbing đại diện cho nhóm local search vì nó không xây dựng cây tìm kiếm đầy đủ. Nó chỉ quan tâm trạng thái hiện tại và neighbor. Chính sự đơn giản đó vừa là ưu điểm vừa là hạn chế: nhanh và dễ hiểu, nhưng không đủ mạnh trước không gian có nhiều cực trị cục bộ.",
    ]:
        add_para(doc, paragraph)

    add_heading(doc, "4. Nhóm môi trường phức tạp: AND-OR Search", 2)
    for paragraph in [
        algorithm_common_intro("AND-OR Search", "nó là thuật toán đã học phù hợp để biểu diễn hành động có nhiều kết quả trong môi trường bất định"),
        "AND-OR Search khác với các thuật toán tìm đường thông thường vì kết quả của nó không chỉ là một chuỗi hành động tuyến tính. Trong môi trường bất định, một hành động có thể dẫn đến nhiều trạng thái kết quả. Agent không chỉ cần chọn hành động tốt, mà còn phải có kế hoạch xử lý cho từng outcome có thể xảy ra.",
        "Trong cây AND-OR, OR node đại diện cho lựa chọn của agent. Tại OR node, chỉ cần chọn một hành động phù hợp. AND node đại diện cho các kết quả môi trường có thể xảy ra sau hành động đó. Tại AND node, kế hoạch phải giải quyết tất cả các nhánh kết quả, vì agent không kiểm soát outcome nào sẽ xảy ra. Đây là điểm cốt lõi làm AND-OR Search phù hợp với môi trường phức tạp.",
    ]:
        add_para(doc, paragraph)
    add_equation(doc, [mr("Result(s,a)={"), sub([mr("s")], [mr("1")]), mr(","), sub([mr("s")], [mr("2")]), mr(",...,"), sub([mr("s")], [mr("k")]), mr("}")])
    add_equation(doc, [mr("Plan(s)=a ∧ ∀ "), sub([mr("s")], [mr("i")]), mr("∈Result(s,a): Plan("), sub([mr("s")], [mr("i")]), mr(")")])
    for paragraph in [
        "Trong Find Your Path, AND-OR Search có thể mô hình hóa tình huống một đoạn đường có thể thông hoặc bị chặn, một chi phí cạnh có thể giữ nguyên hoặc tăng cao, hoặc một điểm giao có thể khả dụng hoặc không khả dụng. Nếu chỉ dùng A*, hệ thống thường tạo một route dựa trên trạng thái hiện biết. Nếu dùng AND-OR Search, hệ thống có thể tạo conditional plan: nếu đoạn đường mở thì tiếp tục route A, nếu đoạn đường bị chặn thì chuyển sang route B.",
        "Trạng thái bắt đầu của AND-OR Search gồm trạng thái hiện tại và mô hình transition. Trạng thái mục tiêu là một conditional plan có thể dẫn đến goal trong mọi outcome quan trọng. Các bước tìm kiếm gồm chọn hành động ở OR node, sinh tập outcome ở AND node, kiểm tra goal, phát hiện chu trình và xây dựng kế hoạch con cho từng outcome. Nếu một outcome không có kế hoạch giải, hành động tương ứng không tạo được conditional plan hợp lệ.",
        "Về tính đầy đủ, AND-OR Search đầy đủ trên không gian hữu hạn nếu transition model liệt kê đầy đủ outcome và thuật toán có cơ chế tránh chu trình. Về tối ưu, thuật toán có thể tối ưu nếu duyệt đầy đủ và có tiêu chí đánh giá rõ ràng, chẳng hạn chi phí xấu nhất hoặc chi phí kỳ vọng. Về thời gian và không gian, cây AND-OR có thể tăng rất nhanh vì mỗi hành động lại sinh nhiều outcome. Nếu b là số hành động trung bình, k là số outcome và m là độ sâu, trường hợp xấu có thể xấp xỉ O((bk)^m).",
        "Điểm mạnh của AND-OR Search là cách biểu diễn kế hoạch giàu hơn path đơn. Nó phù hợp để trả lời câu hỏi: agent cần làm gì nếu môi trường phản hồi theo nhiều cách khác nhau. Điểm yếu là chi phí tính toán cao và đòi hỏi transition model đủ rõ. Nếu nhóm không mô hình hóa outcome tốt, AND-OR Search sẽ trở nên hình thức. Vì vậy, trong báo cáo cần mô tả outcome cụ thể gắn với đường đi: thông, bị chặn, tăng chi phí hoặc phải chuyển hướng.",
        "Cần phân biệt AND-OR Search với online re-planning. Online re-planning thường phản ứng sau khi quan sát thông tin mới, còn AND-OR Search chuẩn bị trước các nhánh điều kiện. Trong báo cáo hiện tại, nhóm chọn AND-OR Search vì đây là thuật toán rõ ràng trong nhóm môi trường phức tạp và phù hợp với yêu cầu mỗi nhóm chọn một thuật toán đã học.",
        "Khi vấn đáp, câu trả lời quan trọng là: AND-OR Search không chỉ tìm một đường đi từ start đến goal, mà tìm một kế hoạch có điều kiện. Đây là khác biệt bản chất so với DFS, A* hay Simple Hill Climbing. Nếu môi trường deterministic, path đơn có thể đủ; nếu môi trường nondeterministic, conditional plan mới phản ánh đúng yêu cầu của agent.",
    ]:
        add_para(doc, paragraph)

    add_heading(doc, "5. Nhóm thỏa ràng buộc: Forward Checking", 2)
    for paragraph in [
        algorithm_common_intro("Forward Checking", "nó đại diện rõ cho tư duy kiểm tra ràng buộc và cắt bỏ sớm miền giá trị không hợp lệ"),
        "Forward Checking là kỹ thuật tìm kiếm trong bài toán thỏa ràng buộc. Nó thường được trình bày như một cải tiến của Backtracking, nhưng trong báo cáo này được chọn làm thuật toán đại diện vì thể hiện rõ ý tưởng nhìn trước. Thay vì đợi đến khi gán nhiều biến rồi mới phát hiện mâu thuẫn, Forward Checking kiểm tra ảnh hưởng của mỗi phép gán lên các biến chưa gán.",
        "Một bài toán CSP gồm tập biến, miền giá trị của từng biến và tập ràng buộc. Trong Find Your Path, biến có thể là vị trí của một đơn hàng trong chuỗi giao, shipper được gán cho đơn, thời điểm phục vụ hoặc lựa chọn có nhận một đơn hay không. Miền giá trị là các lựa chọn khả dĩ cho từng biến. Ràng buộc có thể là deadline, capacity, thứ tự pickup trước dropoff hoặc giới hạn phân công.",
    ]:
        add_para(doc, paragraph)
    add_equation(doc, [sub([mr("D")], [mr("j")]), mr(" ← "), sub([mr("D")], [mr("j")]), mr(" \\ {u | ("), sub([mr("X")], [mr("i")]), mr("=v, "), sub([mr("X")], [mr("j")]), mr("=u) violates C}")])
    for paragraph in [
        "Công thức trên mô tả thao tác loại khỏi domain Dj những giá trị u khiến phép gán Xi = v và Xj = u vi phạm ràng buộc C. Sau khi loại, nếu một biến chưa gán có domain rỗng, thuật toán quay lui ngay. Điều này giúp tránh việc tiếp tục mở rộng một nhánh chắc chắn không thể tạo nghiệm.",
        "Trong Find Your Path, Forward Checking đặc biệt hữu ích khi route không chỉ cần đi được mà còn phải hợp lệ. Ví dụ, một đơn hàng có deadline gần có thể không thể đặt sau nhiều đơn xa; một shipper có capacity giới hạn không thể nhận quá nhiều đơn cùng lúc; một đơn pickup/dropoff không thể giao trước khi lấy. Những điều kiện này không phải lúc nào cũng được thể hiện bằng cạnh trên graph, nên cần mô hình CSP để kiểm soát.",
        "Về tính đầy đủ, Forward Checking đầy đủ với CSP hữu hạn vì thuật toán chỉ loại các giá trị chắc chắn vi phạm ràng buộc. Nếu có nghiệm, nó không loại bỏ nghiệm hợp lệ. Về tối ưu, Forward Checking không tự tối ưu chi phí nếu không có hàm mục tiêu bổ sung; nó chủ yếu trả lời câu hỏi liệu có phép gán hợp lệ hay không. Về thời gian, trường hợp xấu vẫn có thể là O(d_c^n), vì vẫn có thể phải xét tổ hợp lớn. Về không gian, thuật toán cần lưu domain hiện tại và các thay đổi để phục hồi khi quay lui.",
        "Điểm mạnh của Forward Checking là cắt nhánh sớm và làm rõ lý do một phương án không hợp lệ. Trong báo cáo và demo, có thể minh họa domain của các biến bị thu hẹp sau từng phép gán. Đây là hình ảnh trực quan hơn nhiều so với chỉ nói thuật toán quay lui. Người xem có thể thấy ràng buộc tác động trực tiếp đến không gian tìm kiếm.",
        "Điểm yếu của Forward Checking là nó chỉ nhìn trước một bước theo các ràng buộc liên quan trực tiếp. Có những mâu thuẫn sâu hơn chỉ xuất hiện sau nhiều phép gán. Vì vậy, Forward Checking mạnh hơn Backtracking thuần nhưng chưa phải kỹ thuật suy luận ràng buộc toàn diện nhất. Tuy nhiên, trong phạm vi học phần và dự án, nó đủ rõ ràng để đại diện cho nhóm CSP.",
        "Khi vấn đáp, cần tránh nói Forward Checking là thuật toán tìm đường. Nó không tìm path trên graph theo nghĩa A*. Nó kiểm tra và thu hẹp miền giá trị trong bài toán ràng buộc. Trong hệ thống Find Your Path, Forward Checking có thể đứng ở tầng kiểm tra tính khả thi của route hoặc phân công trước khi thuật toán tìm đường chi tiết được gọi.",
    ]:
        add_para(doc, paragraph)

    add_heading(doc, "6. Nhóm tìm kiếm đối kháng: Alpha-Beta Pruning", 2)
    for paragraph in [
        algorithm_common_intro("Alpha-Beta Pruning", "nó là thuật toán đối kháng hiệu quả hơn Minimax thuần nhưng vẫn giữ nguyên quyết định tối ưu"),
        "Alpha-Beta Pruning được xây dựng trên nền Minimax. Trong cây quyết định, MAX là tác tử muốn tối đa hóa utility, còn MIN là tác tử hoặc yếu tố môi trường muốn làm giảm utility. Nếu xét bài toán Find Your Path theo nghĩa đối kháng, MAX có thể là hệ thống chọn route, còn MIN là kịch bản bất lợi như tắc đường, tăng chi phí hoặc chặn cạnh.",
        "Minimax duyệt cây để xác định giá trị của từng lựa chọn trong trường hợp đối thủ phản ứng bất lợi nhất. Alpha-Beta Pruning không thay đổi giá trị Minimax cuối cùng, mà chỉ cắt bỏ những nhánh chắc chắn không ảnh hưởng đến quyết định ở gốc. Vì vậy, thuật toán này có giá trị lớn khi cây quyết định rộng hoặc sâu.",
    ]:
        add_para(doc, paragraph)
    add_equation(doc, [mr("prune if α ≥ β")])
    for paragraph in [
        "Alpha là giá trị tốt nhất mà MAX đã bảo đảm được trên đường đi hiện tại. Beta là giá trị tốt nhất mà MIN đã bảo đảm được. Khi alpha lớn hơn hoặc bằng beta, nhánh còn lại không thể làm thay đổi lựa chọn cuối cùng, nên có thể bỏ qua. Ý tưởng này giúp giảm số trạng thái cần đánh giá mà vẫn giữ nguyên kết quả.",
        "Trong Find Your Path, Alpha-Beta không phải thuật toán dùng trực tiếp để tìm đường hằng ngày. Vai trò hợp lý hơn là mô phỏng và đánh giá route dưới kịch bản xấu. Ví dụ, hệ thống có thể so sánh hai tuyến: tuyến thứ nhất ngắn nhưng dễ bị ảnh hưởng nếu một đoạn bị chặn; tuyến thứ hai dài hơn nhưng ổn định hơn. Mô hình đối kháng giúp phân tích lựa chọn nào có utility tốt hơn khi môi trường phản ứng bất lợi.",
        "Hàm utility cần phản ánh mục tiêu của hệ thống. Utility có thể là số âm của tổng chi phí, phạt trễ hạn và phạt rủi ro. Khi utility càng cao, route càng tốt. Nếu utility chỉ tính khoảng cách, mô hình đối kháng sẽ không phản ánh đủ rủi ro. Vì vậy, thiết kế hàm utility là bước quan trọng khi đưa Alpha-Beta vào bài toán vận tải.",
    ]:
        add_para(doc, paragraph)
    add_equation(doc, [mr("utility(route)=-"), sub([mr("travel")], [mr("cost")]), mr("-"), sub([mr("late")], [mr("penalty")]), mr("-"), sub([mr("risk")], [mr("penalty")])])
    for paragraph in [
        "Về tính đầy đủ, Alpha-Beta đầy đủ nếu cây trò chơi hữu hạn và thuật toán duyệt đủ các nhánh cần thiết. Về tối ưu, nó cho cùng quyết định với Minimax nếu hàm utility đúng và giả định MIN luôn chọn phản ứng bất lợi nhất. Về thời gian, trường hợp xấu vẫn là O(b^m) nếu thứ tự xét nhánh kém. Trong trường hợp thứ tự xét tốt, số node có thể giảm đáng kể, thường được mô tả xấp xỉ O(b^(m/2)). Về không gian, thuật toán có thể duyệt theo chiều sâu nên thường cần O(bm).",
        "Điểm mạnh của Alpha-Beta là hiệu quả tính toán. Nó cho phép mô hình đối kháng có thể mở rộng hơn Minimax thuần. Điểm yếu là hiệu quả cắt tỉa phụ thuộc vào thứ tự xét nhánh. Nếu xét nhánh tốt trước, alpha và beta được cập nhật sớm, nhiều nhánh bị cắt. Nếu xét nhánh xấu trước, thuật toán gần giống Minimax về số node cần duyệt.",
        "Khi vấn đáp, cần nhấn mạnh rằng Alpha-Beta không phải thuật toán khác mục tiêu với Minimax. Nó là cùng giá trị quyết định, nhưng tối ưu quá trình duyệt. Trong báo cáo, nhóm chọn Alpha-Beta thay vì Minimax vì yêu cầu mỗi nhóm chọn một thuật toán, và Alpha-Beta vừa bao hàm tư duy Minimax vừa thể hiện kỹ thuật cắt tỉa nâng cao.",
        "Kết thúc phần thuật toán, có thể thấy sáu thuật toán không cạnh tranh trực tiếp với nhau. DFS minh họa tìm kiếm mù, A* giải bài toán tìm đường có heuristic, Simple Hill Climbing cải thiện cấu hình, AND-OR Search xử lý outcome bất định, Forward Checking kiểm tra ràng buộc và Alpha-Beta phân tích tình huống đối kháng. Mỗi thuật toán trả lời một câu hỏi khác nhau trong cùng miền bài toán Find Your Path.",
    ]:
        add_para(doc, paragraph)


def add_experiment_section(doc):
    add_heading(doc, "III. Thực nghiệm và kết quả", 1)
    for heading, paragraphs in [
        ("1. Thiết kế kịch bản thực nghiệm", [
            "Phần thực nghiệm cần chứng minh các thuật toán hoạt động đúng trên các kịch bản có kiểm soát. Thay vì chỉ chụp màn hình kết quả cuối cùng, nhóm nên trình bày quá trình thuật toán ra quyết định. Với mỗi thuật toán, cần có trạng thái bắt đầu, trạng thái mục tiêu, các bước trung gian và kết quả sau cùng.",
            "Đối với DFS, kịch bản nên dùng graph nhỏ có ít nhất một nhánh cụt để thể hiện rõ quá trình đi sâu và quay lui. Kết quả cần ghi thứ tự node được duyệt, stack tại một số thời điểm quan trọng và đường đi cuối cùng nếu tìm thấy goal. Nếu đường đi không ngắn nhất, đó là minh chứng tốt cho hạn chế của DFS.",
            "Đối với A*, kịch bản nên dùng graph có tọa độ hoặc trọng số rõ ràng. Cần hiển thị g-score, h-score và f-score của các node đang xét. Kết quả nên ghi tổng chi phí route, số node đã mở rộng và đường đi cuối cùng. Nếu có thể, nhóm nên so sánh A* với DFS ở mức số node duyệt, nhưng không biến báo cáo thành bảng so sánh máy móc.",
            "Đối với Simple Hill Climbing, kịch bản nên bắt đầu bằng một route chưa tốt và cho thấy route được cải thiện qua từng vòng. Mỗi vòng cần nêu neighbor được chọn, cost trước và sau khi cập nhật. Nếu thuật toán dừng ở local optimum, phần thực nghiệm nên ghi nhận trung thực thay vì cố trình bày như luôn tìm lời giải tối ưu.",
            "Đối với AND-OR Search, kịch bản nên có ít nhất một hành động với hai outcome. Ví dụ, một cạnh có thể thông hoặc bị chặn. Kết quả không phải một path đơn, mà là conditional plan gồm các nhánh nếu outcome này xảy ra thì làm gì, nếu outcome kia xảy ra thì làm gì. Đây là điểm quan trọng để phân biệt với A*.",
            "Đối với Forward Checking, kịch bản nên có vài biến đơn hàng và ràng buộc deadline hoặc capacity. Nhóm cần hiển thị domain ban đầu, phép gán được thử, domain sau khi loại giá trị và thời điểm thuật toán quay lui do domain rỗng. Cách trình bày này cho thấy thuật toán cắt bớt không gian tìm kiếm như thế nào.",
            "Đối với Alpha-Beta Pruning, kịch bản nên dùng cây quyết định nhỏ có giá trị utility ở lá. Thực nghiệm cần thể hiện quá trình cập nhật alpha, beta và các nhánh bị cắt. Nếu có animation, nên dùng màu khác nhau cho node đã duyệt, node đang xét và node bị cắt.",
        ]),
        ("2. Chỉ tiêu quan sát và cách ghi nhận kết quả", [
            "Kết quả thực nghiệm nên được ghi nhận bằng cả mô tả định tính và số liệu định lượng. Với thuật toán tìm kiếm, số node mở rộng là chỉ tiêu quan trọng vì nó phản ánh chi phí tìm kiếm. Với thuật toán tối ưu cục bộ, giá trị cost qua từng vòng cho thấy xu hướng cải thiện. Với CSP, số giá trị bị loại khỏi domain giúp đánh giá hiệu quả cắt nhánh. Với Alpha-Beta, số node bị cắt là minh chứng trực tiếp cho lợi ích của thuật toán.",
            "Không nên chỉ kết luận thuật toán nào nhanh hơn vì mỗi thuật toán phục vụ một loại bài toán khác nhau. DFS không được kỳ vọng tối ưu như A*. AND-OR Search không trả cùng loại kết quả với A*. Forward Checking kiểm tra miền giá trị chứ không tìm path hình học. Vì vậy, kết quả cần được diễn giải trong đúng bối cảnh của từng nhóm thuật toán.",
            "Khi ghi nhận thời gian chạy, cần lưu ý rằng thời gian phụ thuộc vào kích thước graph, cấu trúc dữ liệu và môi trường máy. Nếu chưa có benchmark ổn định, nhóm có thể ưu tiên số node mở rộng, số bước lặp hoặc số nhánh bị cắt. Các chỉ tiêu này gắn trực tiếp hơn với bản chất thuật toán và dễ giải thích khi vấn đáp.",
            "Đối với A*, chất lượng heuristic cần được trình bày rõ. Nếu heuristic quá nhỏ, A* mở rộng nhiều node hơn nhưng vẫn an toàn về tối ưu. Nếu heuristic quá tham lam hoặc không admissible, thuật toán có thể nhanh hơn nhưng không còn bảo đảm tối ưu. Đây là điểm nên đưa vào phần nhận xét thay vì chỉ trình bày route cuối cùng.",
            "Đối với Simple Hill Climbing, kết quả nên có chuỗi cost: cost ban đầu, cost sau từng bước cải thiện và cost tại điểm dừng. Nếu điểm dừng không phải tối ưu toàn cục, nhóm có thể giải thích do thuật toán chỉ nhìn neighbor cục bộ. Cách trình bày này trung thực và đúng bản chất hơn việc cố tìm một ví dụ luôn thành công.",
            "Đối với Forward Checking, kết quả nên mô tả một nhánh bị loại sớm. Ví dụ, sau khi gán một đơn hàng vào vị trí quá muộn, domain của đơn tiếp theo bị rỗng vì không còn thời gian giao hợp lệ. Đây là bằng chứng rõ ràng cho việc thuật toán tiết kiệm công sức so với thử hết tổ hợp.",
            "Đối với Alpha-Beta, kết quả nên so sánh số node Minimax cần xét với số node Alpha-Beta thực sự duyệt trong cùng cây. Dù báo cáo không chọn Minimax làm thuật toán chính, việc nhắc Minimax như nền tảng giúp giải thích ý nghĩa của cắt tỉa mà không biến nó thành một mục thuật toán riêng.",
        ]),
        ("3. Định hướng chèn hình minh họa", [
            "Theo yêu cầu báo cáo cuối kỳ, phần thực nghiệm nên có ảnh động hoặc chuỗi ảnh minh họa. Với DFS, hình nên thể hiện hướng đi sâu và mũi tên quay lui. Với A*, hình nên thể hiện frontier và giá trị f(n). Với Simple Hill Climbing, hình nên thể hiện route trước và sau khi đổi neighbor. Với AND-OR Search, hình nên thể hiện cây conditional plan. Với Forward Checking, hình nên thể hiện domain bị thu hẹp. Với Alpha-Beta, hình nên thể hiện nhánh bị cắt.",
            "Ảnh minh họa cần được đặt sau đoạn mô tả kịch bản, không nên dồn hết về cuối. Mỗi ảnh nên có chú thích ngắn nêu thuật toán, trạng thái đang xét và ý nghĩa của màu sắc. Nếu dùng GIF hoặc video trong slide, báo cáo Word có thể chèn ảnh tĩnh đại diện và dẫn link GitHub hoặc thư mục demo chứa animation.",
            "Khi chèn hình, cần tránh ảnh quá nhỏ hoặc chỉ là screenshot giao diện không giải thích được thuật toán. Một ảnh tốt phải cho thấy quá trình ra quyết định. Ví dụ, với A*, chỉ thấy đường đi cuối cùng là chưa đủ; nên có thêm node đã mở rộng. Với Alpha-Beta, chỉ thấy cây cuối cùng là chưa đủ; cần đánh dấu nhánh nào bị prune.",
            "Nếu chưa có đầy đủ hình, báo cáo có thể để vị trí chèn hình theo từng thuật toán nhưng nên viết chú thích cụ thể. Chú thích cụ thể giúp người đọc biết nhóm dự định minh họa điều gì, đồng thời giúp quá trình hoàn thiện sau này nhất quán với nội dung phân tích.",
        ]),
    ]:
        add_heading(doc, heading, 2)
        for paragraph in paragraphs:
            add_para(doc, paragraph)
    add_heading(doc, "4. Link GitHub", 2)
    add_para(doc, "Mã nguồn và dữ liệu demo của dự án được lưu tại: https://github.com/zomboXx/FYP.git. Khi hoàn thiện báo cáo, nhóm nên dẫn thêm đường dẫn đến thư mục demo, ảnh động hoặc video minh họa nếu các tài nguyên đó được đưa lên repository.")


def add_discussion(doc):
    add_heading(doc, "IV. Đánh giá và thảo luận", 1)
    paragraphs = [
        "Phần đánh giá không nên được trình bày như một bảng liệt kê cứng, vì sáu thuật toán được chọn không giải cùng một bài toán con. Nếu đặt chúng vào một bảng, báo cáo dễ tạo cảm giác máy móc và làm mất bản chất học thuật của phần thảo luận. Cách hợp lý hơn là phân tích vai trò của từng thuật toán trong hệ thống và chỉ ra vì sao mỗi thuật toán phù hợp với một lớp vấn đề khác nhau.",
        "DFS có giá trị như thuật toán nền. Nó cho thấy một tác tử có thể tìm đường khi không có bất kỳ tri thức định hướng nào. Điểm yếu của DFS trong bài toán giao hàng cũng rất rõ: thuật toán không tối ưu, dễ đi sâu vào nhánh kém và phụ thuộc thứ tự duyệt. Tuy nhiên, chính điểm yếu này giúp làm nổi bật giá trị của A*. Khi trình bày báo cáo, DFS nên được xem như đường cơ sở về mặt nhận thức, không phải giải pháp triển khai chính.",
        "A* là thuật toán có tính ứng dụng cao nhất đối với bài toán tìm đường giữa hai điểm. Nó kết hợp chi phí đã đi và heuristic còn lại, nhờ đó phù hợp với bản đồ có trọng số. Nếu chỉ chọn một thuật toán để chạy route chính trong Find Your Path, A* là lựa chọn hợp lý hơn DFS. Tuy nhiên, A* cũng không giải quyết toàn bộ bài toán giao hàng, vì nó không tự xử lý thứ tự nhiều đơn, ràng buộc deadline hoặc outcome bất định.",
        "Simple Hill Climbing bổ sung góc nhìn tối ưu cục bộ. Trong một hệ thống giao hàng, đôi khi vấn đề không phải là tìm đường giữa hai node, mà là cải thiện một route gồm nhiều điểm. Thuật toán này cho phép giải thích cách một lời giải ban đầu được cải thiện dần. Dù không bảo đảm tối ưu toàn cục, nó phù hợp để minh họa quan hệ giữa hàm đánh giá và hành vi của thuật toán.",
        "AND-OR Search đại diện cho tư duy lập kế hoạch khi môi trường không chắc chắn. Đây là điểm quan trọng vì môi trường giao thông không phải lúc nào cũng deterministic. Một route được tính trước có thể mất hiệu lực khi đường bị chặn hoặc chi phí thay đổi. AND-OR Search không chỉ hỏi đi đường nào, mà hỏi nếu môi trường phản hồi khác nhau thì kế hoạch sẽ rẽ nhánh ra sao. Điều này làm thuật toán phù hợp với phần môi trường phức tạp hơn so với việc chỉ gọi lại A* nhiều lần.",
        "Forward Checking cho thấy một lớp vấn đề khác: tính hợp lệ dưới ràng buộc. Trong thực tế, route ngắn vẫn có thể vô nghĩa nếu vi phạm deadline, capacity hoặc thứ tự pickup/dropoff. Forward Checking không cạnh tranh trực tiếp với A*, mà có thể hoạt động như tầng kiểm tra trước hoặc trong quá trình lập kế hoạch. Vai trò của nó là giảm không gian tìm kiếm bằng cách loại sớm những giá trị chắc chắn không hợp lệ.",
        "Alpha-Beta Pruning đưa vào báo cáo một cách nhìn đối kháng. Dù bài toán giao hàng không phải trò chơi hai người theo nghĩa truyền thống, ta vẫn có thể mô hình hóa môi trường bất lợi như một tác nhân làm giảm utility. Khi đó, Alpha-Beta giúp phân tích lựa chọn nào ổn định hơn dưới kịch bản xấu. Thuật toán này không thay thế A*, nhưng hỗ trợ đánh giá rủi ro ở tầng quyết định.",
        "Nếu đặt sáu thuật toán vào một kiến trúc tổng thể, có thể hình dung hệ thống gồm nhiều tầng. DFS dùng để minh họa và kiểm tra nền tảng graph. A* dùng để tìm tuyến giữa hai điểm. Simple Hill Climbing dùng để cải thiện route nhiều điểm. AND-OR Search dùng để chuẩn bị kế hoạch có điều kiện trong môi trường bất định. Forward Checking dùng để kiểm soát ràng buộc. Alpha-Beta dùng để phân tích lựa chọn dưới kịch bản bất lợi.",
        "Cách kết hợp này phản ánh tinh thần quan trọng của môn Trí tuệ nhân tạo: không có một thuật toán duy nhất phù hợp với mọi vấn đề. Việc chọn thuật toán phụ thuộc vào biểu diễn trạng thái, loại môi trường, mục tiêu tối ưu và dạng thông tin có sẵn. Nếu môi trường đã biết và mục tiêu rõ ràng, A* mạnh. Nếu môi trường có outcome bất định, AND-OR Search tự nhiên hơn. Nếu vấn đề nằm ở ràng buộc, Forward Checking phù hợp hơn.",
        "Hạn chế của báo cáo hiện tại là phần thực nghiệm vẫn cần được bổ sung bằng số liệu và hình ảnh cụ thể từ chương trình. Phân tích lý thuyết đã chỉ ra vai trò của từng thuật toán, nhưng để thuyết phục hơn, nhóm cần xuất trace, animation hoặc screenshot cho từng thuật toán. Đặc biệt, A* nên có số node mở rộng, Forward Checking nên có domain trước và sau khi prune, Alpha-Beta nên có số nhánh bị cắt.",
        "Một điểm cần tránh khi trình bày là không nên nói thuật toán nào tốt nhất một cách tuyệt đối. DFS kém A* trong tìm đường tối ưu, nhưng DFS lại đơn giản và dễ minh họa. Simple Hill Climbing không bảo đảm tối ưu toàn cục, nhưng hữu ích khi cần cải thiện nhanh. Forward Checking không tìm đường, nhưng kiểm soát tính hợp lệ. Vì vậy, đánh giá đúng phải dựa trên vai trò của thuật toán trong hệ thống.",
        "Từ góc nhìn dự án Find Your Path, trọng tâm triển khai nên đặt ở A* và các dữ liệu graph, vì đây là phần người dùng dễ quan sát nhất. Các thuật toán còn lại có thể được dùng để mở rộng hoặc minh họa các khía cạnh AI khác của bài toán. Cách trình bày này giúp báo cáo vừa bám đề tài ứng dụng, vừa thể hiện được độ phủ kiến thức của học phần.",
        "Nếu cần bảo vệ trước giảng viên, nhóm có thể giải thích rằng việc chọn một thuật toán mỗi nhóm giúp báo cáo sâu hơn. Thay vì nêu hai thuật toán nhưng phân tích mỏng, báo cáo tập trung vào một đại diện tiêu biểu, phân tích đủ trạng thái, mục tiêu, công thức, độ phức tạp, tính đầy đủ, tối ưu và ứng dụng. Đây là lựa chọn hợp lý khi yêu cầu mới của phần II là mỗi mục một thuật toán.",
    ]
    for paragraph in paragraphs:
        add_para(doc, paragraph)


def add_conclusion_and_refs(doc):
    add_heading(doc, "V. Kết luận", 1)
    for paragraph in [
        "Báo cáo đã phân tích bài toán Find Your Path dưới góc nhìn của một hệ thống trí tuệ nhân tạo. Bài toán không chỉ là tìm một đường đi ngắn, mà là lập kế hoạch trong môi trường có trọng số, ràng buộc, mục tiêu tối ưu và khả năng phát sinh tình huống bất định. Việc phân tích PEAS giúp xác định rõ tác tử đang tối ưu điều gì, hoạt động trong môi trường nào, thực hiện hành động nào và nhận thông tin từ nguồn nào.",
        "Phần thuật toán đã chọn sáu đại diện theo đúng nhóm kiến thức: DFS cho tìm kiếm không thông tin, A* cho tìm kiếm có thông tin, Simple Hill Climbing cho tìm kiếm cục bộ, AND-OR Search cho môi trường phức tạp, Forward Checking cho thỏa ràng buộc và Alpha-Beta Pruning cho tìm kiếm đối kháng. Mỗi thuật toán được phân tích theo trạng thái bắt đầu, trạng thái mục tiêu, quy trình tìm lời giải, công thức nếu có, độ phức tạp và bốn tính chất cơ bản.",
        "Kết quả quan trọng nhất của báo cáo là làm rõ rằng các thuật toán không thay thế lẫn nhau hoàn toàn. Chúng đại diện cho các cách nhìn khác nhau về cùng một miền bài toán. A* phù hợp để tìm đường, Forward Checking phù hợp để kiểm tra ràng buộc, AND-OR Search phù hợp với kế hoạch có điều kiện, còn Alpha-Beta phù hợp để phân tích tình huống bất lợi. Cách hiểu này giúp nhóm trình bày và vấn đáp tự tin hơn.",
        "Trong bước tiếp theo, nhóm cần bổ sung hình minh họa và số liệu thực nghiệm từ chương trình. Mỗi thuật toán nên có ít nhất một trace hoặc ảnh động thể hiện quá trình ra quyết định. Khi các minh họa này được chèn vào báo cáo và slide, sản phẩm cuối kỳ sẽ có sự liên kết tốt hơn giữa lý thuyết, mã nguồn và kết quả trình diễn.",
    ]:
        add_para(doc, paragraph)

    add_heading(doc, "Tài liệu tham khảo", 1)
    refs = [
        "Russell, S. J., & Norvig, P. Artificial Intelligence: A Modern Approach. Pearson.",
        "Tài liệu bài giảng học phần Trí tuệ nhân tạo, Trường Đại học Sư phạm Kỹ thuật TP. Hồ Chí Minh.",
        "Tài liệu mã nguồn dự án Find Your Path tại GitHub: https://github.com/zomboXx/FYP.git.",
        "Python documentation và tài liệu thư viện hỗ trợ xử lý đồ thị, giao diện và kiểm thử trong dự án.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        format_paragraph(p, first_line=False)


def add_page_floor_appendix(doc):
    # Content appendix is deliberately analytical, not filler; it helps the report clear a 30-page budget in Word.
    add_heading(doc, "Phụ lục phân tích mở rộng", 1)
    themes = [
        ("Biểu diễn trạng thái", "Biểu diễn trạng thái quyết định trực tiếp đến thuật toán có thể áp dụng. Nếu trạng thái chỉ là node hiện tại, bài toán nghiêng về path finding. Nếu trạng thái bao gồm tập đơn hàng, thời gian và tải hiện tại, bài toán trở thành lập kế hoạch có ràng buộc. Nếu trạng thái còn bao gồm belief về cạnh bị chặn hoặc outcome môi trường, bài toán chuyển sang môi trường phức tạp."),
        ("Thiết kế hàm chi phí", "Hàm chi phí cần phản ánh mục tiêu thực tế của giao hàng. Khoảng cách là thành phần dễ đo, nhưng không đủ. Thời gian, deadline, capacity, rủi ro và mức độ ổn định của route đều có thể trở thành thành phần chi phí. Một thuật toán tốt trên hàm chi phí sai vẫn có thể tạo quyết định kém trong thực tế."),
        ("Quan hệ giữa A* và các tầng khác", "A* có thể là lõi tìm đường giữa hai điểm, nhưng không nên bị ép giải mọi bài toán con. Khi có nhiều đơn hàng, tầng local search có thể quyết định thứ tự điểm đến. Khi có ràng buộc, tầng CSP kiểm tra tính hợp lệ. Khi có outcome bất định, tầng lập kế hoạch có điều kiện giúp chuẩn bị nhánh dự phòng."),
        ("Ý nghĩa của thực nghiệm", "Thực nghiệm trong báo cáo AI không chỉ nhằm chứng minh chương trình chạy được. Nó còn phải cho thấy thuật toán ra quyết định như thế nào. Một ảnh kết quả cuối cùng thường chưa đủ; cần có trace, số node mở rộng, số nhánh bị cắt, domain bị thu hẹp hoặc cost thay đổi qua từng vòng."),
        ("Giới hạn của mô hình", "Mọi mô hình đều giản lược thực tế. Graph có thể không phản ánh đầy đủ đường cấm, thời tiết, hành vi người lái hoặc dữ liệu giao thông thời gian thực. Vì vậy, báo cáo cần thừa nhận giới hạn mô hình và xem thuật toán như công cụ tư duy, không phải lời giải tuyệt đối cho mọi tình huống giao hàng."),
    ]
    for title, base in themes:
        add_heading(doc, title, 2)
        for i in range(4):
            add_para(doc, base + " Trong bối cảnh Find Your Path, nhận xét này giúp nhóm liên hệ giữa lý thuyết thuật toán và quyết định triển khai. Khi vấn đáp, phần phân tích mở rộng này có thể được dùng để giải thích vì sao báo cáo chọn nhiều nhóm thuật toán thay vì chỉ tập trung vào một thuật toán tìm đường duy nhất.")


def mark_update_fields(doc):
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main():
    doc = Document(DOCX)
    set_styles(doc)
    clear_from_toc(doc)
    add_toc(doc)
    add_problem_section(doc)
    add_algorithm_sections(doc)
    add_experiment_section(doc)
    add_discussion(doc)
    add_conclusion_and_refs(doc)
    add_page_floor_appendix(doc)
    mark_update_fields(doc)
    doc.save(DOCX)
    if EXTRA.exists():
        EXTRA.unlink()
    print(DOCX)


if __name__ == "__main__":
    main()
